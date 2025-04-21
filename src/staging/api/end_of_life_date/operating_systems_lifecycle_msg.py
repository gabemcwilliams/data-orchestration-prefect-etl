# Windows Lifecycle Update Report Generator (Scrubbed)

import pandas as pd
import re
import os
import datetime as dt
import requests
from docx import Document

export_dir = "d:/exports"  # Change this to your desired export path


class WindowsBuildResults:
    def __init__(self):
        self.__df = self.__get_endoflife__()

    @staticmethod
    def __get_endoflife__():
        resp = requests.get("https://endoflife.date/api/windows.json")
        df = pd.read_json(resp.text)
        df[["releaseDate", "support", "eol"]] = df[["releaseDate", "support", "eol"]].apply(pd.to_datetime)
        return df

    def eol_dict(self):
        win_os_nums = []
        min_build_list = []
        max_build_list = []
        df_dict = {
            "eol": self.__df[self.__df["support"] < dt.datetime.now()].head(10).sort_values(by="support", ascending=False),
            "newest": self.__df[self.__df["support"] > dt.datetime.now()].head(10).sort_values(by="releaseDate", ascending=False)
        }

        for k, v in df_dict.items():
            for index, row in v.iterrows():
                result = re.match(r"(\d{2}),\\sversion\\s(\w{4})\\s?([\(\)EWLTSIoT]+)\\s?", row["cycle"])
                if result:
                    win_num = result.group(1)
                    win_edi = (result.group(3)).strip()
                    if win_num not in win_os_nums:
                        if win_edi in ["(W)", "(E)(W)"]:
                            win_os_nums.append(result.group(1))
                            build_entry = {
                                "win_num": win_num,
                                "release_name": result.group(2),
                                "release_date": row["releaseDate"],
                                "end_of_support_date": row["support"],
                                "edition": win_edi,
                                "is_lts": row["lts"],
                                "link": row["link"]
                            }
                            if k == "newest":
                                max_build_list.append(build_entry)
                            else:
                                min_build_list.append(build_entry)

        return {
            "win_os_nums": win_os_nums,
            "newest": {"build_list": max_build_list, "df": df_dict["newest"]},
            "eol": {"build_list": min_build_list, "df": df_dict["eol"]}
        }


def add_suffix_to_day(day):
    day = int(day)
    suffix = {1: 'st', 2: 'nd', 3: 'rd', 21: 'st', 22: 'nd', 23: 'rd', 31: 'st'}.get(day, 'th')
    num = re.sub(r"0\d", "", str(day))
    return num + suffix + ","


# Initialization and values
windows_builds = WindowsBuildResults()
results = windows_builds.eol_dict()
windows_os = results["win_os_nums"]

if len(windows_os) == 1:
    os_desc = f"Windows {windows_os[0]}"
    s = ""
elif len(windows_os) == 2:
    os_desc = f"Windows {windows_os[0]} and {windows_os[1]}"
    s = "s"
else:
    exit(1)

newest_version = results["newest"]["build_list"][0]["release_name"]
newest_version_date = " ".join([
    results["newest"]["build_list"][0]["release_date"].month_name(),
    add_suffix_to_day(results["newest"]["build_list"][0]["release_date"].day),
    str(results["newest"]["build_list"][0]["release_date"].year)
])

min_end_of_life_version = results["eol"]["build_list"][0]["release_name"]
min_end_of_life_version_date = " ".join([
    results["eol"]["build_list"][0]["end_of_support_date"].month_name(),
    add_suffix_to_day(results["eol"]["build_list"][0]["end_of_support_date"].day),
    str(results["eol"]["build_list"][0]["end_of_support_date"].year)
])

when_we_start_updating_time = "11 PM"
when_we_start_updating_date = " ".join([
    dt.datetime.now().strftime("%B"),
    add_suffix_to_day((dt.datetime.now() + dt.timedelta(7)).strftime("%d")),
    dt.datetime.now().strftime("%Y")
])

# Paragraph builders
def message_greeting(doc):
    p = doc.add_paragraph(f"""
Hi there,

As you may know, Microsoft feature updates for {os_desc} build{s} are released annually. The current release{s} is build {newest_version} ({newest_version_date}). Each build receives quality updates for 18 to 30 months.
""")
    p.add_run(f"\n\nThe latest build to reach end of life is {os_desc} build {min_end_of_life_version}, which will reach end of life on {min_end_of_life_version_date}.\n").bold = True
    return doc

def what_this_means(doc):
    p = doc.add_paragraph()
    p.add_run("\nWhat This Means:").bold = True
    doc.add_paragraph("""
Devices on these builds will no longer receive Microsoft support or updates, including automated patching via your company's services. We recommend upgrading systems still on this version.
""")
    return doc

def what_happens_next(doc):
    p = doc.add_paragraph()
    p.add_run("\nWhat Happens Next:").bold = True
    doc.add_paragraph(f"""
Your company will begin updating all {os_desc} machines at or below build {min_end_of_life_version} starting {when_we_start_updating_date} at {when_we_start_updating_time} EST. Updates occur during the normal maintenance window for clients subscribed to our patching service.
""")
    return doc

def what_should_i_look_out_for(doc):
    items = [
        f"Ensure all {os_desc} devices are online during maintenance windows.",
        "In low-bandwidth environments, large build downloads may cause temporary slowdowns.",
        f"For compatibility concerns, verify software support for build {min_end_of_life_version}."
    ]
    p = doc.add_paragraph()
    p.add_run("\nWhat Should I Look Out For?\n").bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    return doc

def what_if_problems(doc):
    p = doc.add_paragraph()
    p.add_run("\nWhat If Problems Should Arise?").bold = True
    doc.add_paragraph(f"""
If any compatibility or technical issues arise, our team can roll back the update using a tested script. Please contact support@example.com if assistance is needed.
""")
    return doc

def if_you_have_more_questions(doc):
    p = doc.add_paragraph()
    p.add_run("\nNeed something else?").bold = True
    doc.add_paragraph("""
If you have any questions, contact support@example.com.

Thank you,  
Your Company Name Team
""").italic = True
    return doc

def for_more_information(doc):
    p = doc.add_paragraph()
    p.add_run("\nFor More Information:").bold = True
    for i in windows_os:
        policy_url = f"https://docs.microsoft.com/en-us/lifecycle/products/windows-{i}-home-and-pro"
        try:
            resp = requests.get(policy_url)
            if resp.status_code == 200:
                doc.add_paragraph(f'Microsoft Statement on Windows {i} Lifecycle:\n{policy_url}\n')
        except Exception as e:
            print(e)
    return doc

# Generate document
document = Document()
document = message_greeting(document)
document = what_this_means(document)
document = what_happens_next(document)
document.add_page_break()
document = what_should_i_look_out_for(document)
document = what_if_problems(document)
document = if_you_have_more_questions(document)
document.add_page_break()
document = for_more_information(document)

document.save(f"{export_dir}/operating_system_lifecycle.docx")
